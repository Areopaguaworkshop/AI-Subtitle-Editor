from whisper_cpp_python import Whisper
from docx import Document

# Initialize whispercpp in python 
model = Whisper(model_path="/home/ajiap/ai/whisper.cpp/models/ggml-large-v3-turbo.bin")


# Path to the audio file
audio_path = "/home/ajiap/ai/ctcfol-syriac22.m4a"

# Transcribe the audio without timestamps
result = model.transcribe(audio_path, language="zh")

# Extract transcription text (without timestamps)
transcription = result["text"]

# Save as plaintext file without timestamps
with open("transcription.txt", "w") as text_file:
    text_file.write(transcription)

# Save as markdown file without timestamps
with open("transcription.md", "w") as markdown_file:
    markdown_file.write("# Transcription\n\n")
    markdown_file.write(transcription)

# Save as docx file without timestamps
doc = Document()
doc.add_heading('Transcription', 0)
doc.add_paragraph(transcription)
doc.save("transcription.docx")

print("Transcription completed and saved as plaintext, markdown, and docx without timestamps.")
